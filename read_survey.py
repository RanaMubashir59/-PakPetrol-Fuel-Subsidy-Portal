from docx import Document
import json

doc = Document(r'C:\Users\RanaM\Downloads\Statistical survey.docx')

print('=== PARAGRAPHS ===')
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

print('\n=== TABLES ===')
for i, table in enumerate(doc.tables):
    print(f'\nTable {i+1}:')
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(' | '.join(cells))
